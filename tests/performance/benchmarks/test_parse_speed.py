"""
Benchmark: resume text extraction speed.

Tests how fast extract_text() processes PDF and DOCX files.
- DOCX test: always runs, builds a file in memory.
- PDF test:  runs only when sample PDFs exist in tests/data/resumes/.

Run:
    pytest tests/performance/benchmarks/test_parse_speed.py -v
"""

from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document

from backend import extract_text

SAMPLE_RESUMES_DIR = Path(__file__).resolve().parents[3] / "tests" / "data" / "resumes"


# ------------------------------------------------------------------ #
# DOCX benchmark — always runs, no external files needed             #
# ------------------------------------------------------------------ #
def _make_docx_bytes() -> bytes:
    """Build a minimal DOCX in memory without touching the filesystem."""
    doc = Document()
    doc.add_paragraph("Jane Smith")
    doc.add_paragraph("Senior Python Engineer — 8 years FastAPI, PostgreSQL, Docker.")
    doc.add_paragraph("Skills: Python, SQL, Docker, Kubernetes, REST APIs.")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


_DOCX_BYTES = _make_docx_bytes()


def test_docx_parse_speed(benchmark):
    """Benchmark DOCX text extraction. Always runs — no sample files needed."""

    def parse():
        buf = io.BytesIO(_DOCX_BYTES)
        buf.name = "benchmark_resume.docx"
        return extract_text(buf)

    result = benchmark(parse)
    assert result is not None
    assert len(result.strip()) > 0, "DOCX extraction returned empty text"
    assert "Jane Smith" in result


# ------------------------------------------------------------------ #
# PDF benchmark — only runs when sample PDFs exist                   #
# ------------------------------------------------------------------ #
def _pdf_files():
    if not SAMPLE_RESUMES_DIR.is_dir():
        return []
    return sorted(SAMPLE_RESUMES_DIR.glob("*.pdf"))


@pytest.mark.parametrize("resume_path", _pdf_files(), ids=lambda p: p.stem)
def test_pdf_parse_speed(benchmark, resume_path):
    """Benchmark PDF text extraction against real sample resumes."""
    pdf_bytes = resume_path.read_bytes()

    def parse():
        buf = io.BytesIO(pdf_bytes)
        buf.name = resume_path.name
        return extract_text(buf)

    result = benchmark(parse)
    assert result is not None
    assert isinstance(result, str)
